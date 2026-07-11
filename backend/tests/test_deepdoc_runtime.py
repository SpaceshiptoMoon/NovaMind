import sys
import asyncio
from pathlib import Path
from types import SimpleNamespace

import numpy as np
import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from PIL import Image
from PyPDF2 import PdfReader
from io import BytesIO
import zipfile

BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from src.features.knowledge_space.schemas.knowledge_base_schema import KnowledgeBaseConfig
from src.features.knowledge_space.services.knowledge_base_service import KnowledgeBaseService
from src.features.knowledge_space.services.document_service import (
    _extract_parse_metadata_summary,
    _prepare_es_chunks_static,
)
from src.shared.integrations.deepdoc.core.capabilities import get_deepdoc_capabilities
from src.shared.integrations.deepdoc.compat.compat import LazyImage
from src.shared.integrations.deepdoc.diagnostics.dependencies import get_deepdoc_runtime_report
from src.shared.integrations.deepdoc.core.engine import DeepDocEngine
from src.shared.integrations.deepdoc.core.factory import DeepDocParserFactory
from src.shared.utils.deepdoc.page_filter import PageNoiseFilter
from src.shared.utils.deepdoc.parser import DocxParser as UpstreamDocxParserAlias
from src.shared.utils.deepdoc.parser import PdfParser as UpstreamPdfParserAlias
from src.shared.utils.deepdoc.pdf_artifacts import PdfArtifactExtractor
from src.shared.utils.deepdoc.ragflow_pdf_parser import DeepDocPdfBox, RAGFlowPdfParser
from src.shared.utils.deepdoc.parser import DeepDocParser
from src.shared.integrations.deepdoc.core.models import DeepDocParseResult
from src.shared.utils.deepdoc.server import create_deepdoc_app
from src.shared.utils.deepdoc.text_concat_model import get_text_concat_model_status
from src.shared.integrations.deepdoc.compat.upstream import get_upstream_deepdoc_snapshot
from src.shared.utils.deepdoc.updown_concat import UpDownConcatMerger
from src.shared.utils.deepdoc.vision.model_manager import (
    ensure_model_group_available,
    expected_model_files,
    get_model_status,
)
from src.shared.utils.deepdoc.vision.package_status import get_vendored_vision_package_status
from src.shared.utils.deepdoc.vision_runtime import (
    DeepDocVisionParserUnavailable,
    DeepDocVisionRuntimeUnavailable,
    ensure_vision_parser_available,
    ensure_vision_runtime_available,
    get_vision_health_status,
    get_vision_runtime_status,
    run_vision_smoke_check,
)
from src.shared.document_processing.pipeline import DocumentProcessor


def _run(coro):
    return asyncio.run(coro)


def _skip_if_vision_runtime_unavailable():
    status = get_vision_runtime_status()
    if not status["available"]:
        pytest.skip(f"DeepDoc vision runtime unavailable: missing {', '.join(status['missing_required'])}")
    return status


def _build_minimal_pdf_bytes(text: str) -> bytes:
    stream = f"BT\n/F1 18 Tf\n72 100 Td\n({text}) Tj\nET".encode("latin-1")
    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        (
            b"3 0 obj\n"
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\n"
            b"endobj\n"
        ),
        b"4 0 obj\n<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream\nendobj\n",
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
    ]

    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    offsets = [0]
    body = b""
    current = len(header)
    for obj in objects:
        offsets.append(current)
        body += obj
        current += len(obj)

    xref_start = len(header) + len(body)
    xref = [b"xref\n0 6\n", b"0000000000 65535 f \n"]
    for offset in offsets[1:]:
        xref.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    trailer = (
        b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n"
        + str(xref_start).encode("ascii")
        + b"\n%%EOF\n"
    )
    return header + body + b"".join(xref) + trailer


def _build_minimal_png_bytes() -> bytes:
    image = Image.new("RGB", (32, 24), color="white")
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _build_positioned_pdf_bytes(operations: list[tuple[int, int, str]]) -> bytes:
    commands = ["BT", "/F1 14 Tf"]
    for x, y, text in operations:
        escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        commands.extend([f"1 0 0 1 {x} {y} Tm", f"({escaped}) Tj"])
    commands.append("ET")
    stream = "\n".join(commands).encode("latin-1")

    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        (
            b"3 0 obj\n"
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 600 800] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\n"
            b"endobj\n"
        ),
        b"4 0 obj\n<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream\nendobj\n",
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n",
    ]

    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    offsets = [0]
    body = b""
    current = len(header)
    for obj in objects:
        offsets.append(current)
        body += obj
        current += len(obj)

    xref_start = len(header) + len(body)
    xref = [b"xref\n0 6\n", b"0000000000 65535 f \n"]
    for offset in offsets[1:]:
        xref.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    trailer = (
        b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n"
        + str(xref_start).encode("ascii")
        + b"\n%%EOF\n"
    )
    return header + body + b"".join(xref) + trailer


def _build_minimal_epub_bytes(title: str, body: str) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w") as zf:
        zf.writestr("mimetype", "application/epub+zip")
        zf.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>""",
        )
        zf.writestr(
            "OEBPS/content.opf",
            f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="2.0" unique-identifier="BookId">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>{title}</dc:title>
  </metadata>
  <manifest>
    <item id="chapter1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine toc="ncx">
    <itemref idref="chapter1"/>
  </spine>
</package>""",
        )
        zf.writestr(
            "OEBPS/chapter1.xhtml",
            f"""<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <body>
    <h1>{title}</h1>
    <p>{body}</p>
  </body>
</html>""",
        )
    return buffer.getvalue()


def test_knowledge_base_config_accepts_deepdoc_strategy():
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_pdf_mode": "plain"},
            "splitting": {"strategy": "recursive"},
        }
    )

    assert config.model_dump()["parsing"]["strategy"] == "deepdoc"
    assert config.model_dump()["parsing"]["deepdoc_pdf_mode"] == "plain"


def test_knowledge_base_service_accepts_all_deepdoc_vision_options():
    service = object.__new__(KnowledgeBaseService)

    service._validate_config_updates(
        {
            "parsing": {
                "strategy": "deepdoc",
                "deepdoc_parser_id": "pdf_paddleocr",
                "deepdoc_pdf_mode": "vision",
            }
        }
    )


def test_deepdoc_runtime_parser_can_be_constructed_without_optional_format_imports():
    parser = DeepDocParser()
    assert parser is not None
    assert "pdf" in parser.supported_extensions()


def test_deepdoc_package_lazy_exports_do_not_force_excel_or_ppt_imports():
    from src.shared.utils.deepdoc import DeepDocParseResult
    from src.shared.utils.deepdoc.parser import TxtParser

    assert DeepDocParseResult.__name__ == "DeepDocParseResult"
    assert TxtParser.__name__ == "RAGFlowTxtParser"


def test_deepdoc_pdf_parser_can_be_imported_without_vision_or_xgboost_runtime():
    assert RAGFlowPdfParser.__name__ == "RAGFlowPdfParser"


def test_deepdoc_engine_can_be_constructed_without_optional_heavy_runtime_imports():
    engine = DeepDocEngine()
    assert engine is not None
    assert "pdf" in engine.supported_extensions()


def test_document_processor_uses_deepdoc_strategy(monkeypatch, tmp_path):
    doc_path = tmp_path / "sample.pdf"
    doc_path.write_bytes(b"%PDF-1.4\n")

    processor = DocumentProcessor()

    async def fake_parse(file_path, *, parsing_config=None, splitting_config=None):
        assert Path(file_path) == doc_path
        assert parsing_config["strategy"] == "deepdoc"
        assert splitting_config["chunk_size"] == 256
        return DeepDocParseResult(
            full_text="@@1\t0\t10\t0\t10##hello",
            chunks=["@@1\t0\t10\t0\t10##hello"],
            metadata={"parser": "deepdoc"},
        )

    monkeypatch.setattr(processor, "_deepdoc_parser", SimpleNamespace(parse=fake_parse))

    full_text, chunks = _run(
        processor.parse_document(
            doc_path,
            parsing_config={"strategy": "deepdoc"},
            splitting_config={"chunk_size": 256},
        )
    )

    assert full_text.startswith("@@1")
    assert chunks == ["@@1\t0\t10\t0\t10##hello"]


def test_document_processor_uses_deepdoc_parser_id(monkeypatch, tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(_build_minimal_pdf_bytes("Processor Parser ID"))

    processor = DocumentProcessor()

    async def fake_aparse_with_parser_id(**kwargs):
        assert kwargs["file_type"] == "pdf"
        assert kwargs["parser_id"] == "pdf_plain"
        assert Path(kwargs["file_path"]) == pdf_path
        return DeepDocParseResult(
            full_text="processor parser id",
            chunks=["processor parser id"],
            metadata={"parser": "deepdoc", "pdf_mode": "plain"},
        )

    monkeypatch.setattr(processor, "_deepdoc_engine", SimpleNamespace(aparse_with_parser_id=fake_aparse_with_parser_id))

    full_text, chunks = _run(
        processor.parse_document(
            pdf_path,
            parsing_config={"strategy": "deepdoc", "deepdoc_parser_id": "pdf_plain"},
            splitting_config={"chunk_size": 256},
        )
    )

    assert full_text == "processor parser id"
    assert chunks == ["processor parser id"]


def test_document_processor_runs_real_deepdoc_pdf_path(tmp_path):
    pdf_path = tmp_path / "real-deepdoc.pdf"
    pdf_path.write_bytes(_build_minimal_pdf_bytes("Real Processor DeepDoc"))

    processor = DocumentProcessor()
    result = _run(
        processor.parse_document_result(
            pdf_path,
            parsing_config={
                "strategy": "deepdoc",
                "deepdoc_parser_id": "pdf_plain",
            },
            splitting_config={"chunk_size": 256},
        )
    )

    assert "Real Processor DeepDoc" in result.full_text
    assert result.chunks
    assert result.metadata["parser"] == "deepdoc"
    assert result.metadata["parser_id"] == "pdf_plain"


def test_document_processor_parse_document_result_preserves_deepdoc_metadata(monkeypatch, tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(_build_minimal_pdf_bytes("Structured Chunk"))

    processor = DocumentProcessor()

    async def fake_parse(file_path, *, parsing_config=None, splitting_config=None):
        return DeepDocParseResult(
            full_text="structured full text",
            chunks=["chunk 1"],
            metadata={"parser": "deepdoc", "chunk_structure": [{"entry_kinds": ["text"], "pages": [1]}]},
        )

    monkeypatch.setattr(processor, "_deepdoc_parser", SimpleNamespace(parse=fake_parse))

    result = _run(
        processor.parse_document_result(
            pdf_path,
            parsing_config={"strategy": "deepdoc"},
            splitting_config={"chunk_size": 256},
        )
    )

    assert result.full_text == "structured full text"
    assert result.metadata["parser"] == "deepdoc"
    assert result.metadata["chunk_structure"][0]["entry_kinds"] == ["text"]


def test_deepdoc_parser_uses_ragflow_adapted_docx_parser(tmp_path):
    doc_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_heading("System Design", level=1)
    doc.add_paragraph("This is the overview paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Owner"
    table.cell(1, 1).text = "NovaMind"
    doc.save(doc_path)

    parser = DeepDocParser()
    result = _run(parser.parse(doc_path, splitting_config={"chunk_size": 500}))

    assert result.metadata["source"] == "ragflow-adapted"
    assert result.metadata["file_type"] == "docx"
    assert "# System Design" in result.full_text
    assert "overview paragraph" in result.full_text
    assert "<table>" in result.full_text


def test_deepdoc_parser_parses_real_pdf(tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    pdf_bytes = _build_minimal_pdf_bytes("Hello DeepDoc")
    pdf_path.write_bytes(pdf_bytes)

    # sanity check the generated PDF is readable before testing parser behavior
    reader = PdfReader(str(pdf_path))
    assert "Hello DeepDoc" in (reader.pages[0].extract_text() or "")

    parser = DeepDocParser()
    result = _run(parser.parse(pdf_path, splitting_config={"chunk_size": 500}))

    assert result.metadata["source"] == "ragflow-adapted"
    assert result.metadata["file_type"] == "pdf"
    assert "@@1\t" in result.full_text
    assert "Hello DeepDoc" in result.full_text
    assert result.chunks
    assert result.metadata["parser_class"] == "RAGFlowPdfParser"
    assert result.metadata["paragraph_merge_strategy"] in {"heuristic", "xgboost"}
    assert "text_concat_model" in result.metadata


def test_deepdoc_parser_supports_textish_formats(tmp_path):
    txt_path = tmp_path / "notes.md"
    txt_path.write_text("# Intro\nDeepDoc note", encoding="utf-8")

    parser = DeepDocParser()
    result = _run(parser.parse(txt_path, splitting_config={"chunk_size": 500}))

    assert result.metadata["file_type"] == "md"
    assert result.metadata["source"] == "ragflow-adapted"
    assert result.metadata["parser_class"] == "RAGFlowMarkdownParser"
    assert "DeepDoc note" in result.full_text
    assert result.chunks


def test_deepdoc_parser_uses_ragflow_html_parser():
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            b"<html><body><h1>Title</h1><p>Hello <b>DeepDoc</b></p><table><tr><td>A</td></tr></table></body></html>",
            file_type="html",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "html"
    assert result.metadata["parser_class"] == "RAGFlowHtmlParser"
    assert "Title" in result.full_text
    assert "<table>" in result.full_text


def test_deepdoc_parser_uses_ragflow_json_parser():
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            b'{"title":"DeepDoc","items":[{"name":"alpha"},{"name":"beta"}]}',
            file_type="json",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "json"
    assert result.metadata["parser_class"] == "RAGFlowJsonParser"
    assert '"title":"DeepDoc"' in result.full_text
    assert result.chunks


def test_deepdoc_parser_supports_parse_bytes_for_text():
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            b"# Intro\nByte mode DeepDoc",
            file_type="md",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "md"
    assert "Byte mode DeepDoc" in result.full_text
    assert result.chunks


def test_deepdoc_parser_supports_parse_bytes_for_docx():
    buffer = BytesIO()
    doc = DocxDocument()
    doc.add_heading("Bytes Doc", level=1)
    doc.add_paragraph("Docx bytes path")
    doc.save(buffer)

    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            buffer.getvalue(),
            file_type="docx",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "docx"
    assert "Bytes Doc" in result.full_text
    assert result.chunks


def test_deepdoc_parser_supports_parse_bytes_for_epub():
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_epub_bytes("DeepDoc EPUB", "Chapter body"),
            file_type="epub",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "epub"
    assert result.metadata["parser_class"] == "RAGFlowEpubParser"
    assert "DeepDoc EPUB" in result.full_text
    assert "Chapter body" in result.full_text
    assert result.chunks


def test_deepdoc_parser_supports_parse_bytes_for_excel():
    Workbook = pytest.importorskip("openpyxl").Workbook
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Metrics"
    sheet.append(["Name", "Value"])
    sheet.append(["Accuracy", "0.92"])
    buffer = BytesIO()
    workbook.save(buffer)

    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            buffer.getvalue(),
            file_type="xlsx",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "xlsx"
    assert result.metadata["parser_class"] == "RAGFlowExcelParser"
    assert "<table>" in result.full_text
    assert "Accuracy" in result.full_text
    assert result.chunks


def test_deepdoc_parser_supports_parse_bytes_for_ppt():
    Presentation = pytest.importorskip("pptx").Presentation
    presentation = Presentation()
    slide = presentation.slides.add_slide(presentation.slide_layouts[1])
    slide.shapes.title.text = "DeepDoc Deck"
    slide.placeholders[1].text = "Slide body"
    buffer = BytesIO()
    presentation.save(buffer)

    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            buffer.getvalue(),
            file_type="pptx",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "pptx"
    assert result.metadata["parser_class"] == "RAGFlowPptParser"
    assert "DeepDoc Deck" in result.full_text
    assert "# Slide 1" in result.full_text
    assert result.chunks


def test_deepdoc_parser_supports_parse_bytes_for_figure():
    _skip_if_vision_runtime_unavailable()
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_png_bytes(),
            file_type="png",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "png"
    assert result.metadata["parser_class"] == "RAGFlowFigureParser"
    assert result.metadata["image"]["width"] == 32
    assert "Image file" in result.full_text
    assert result.chunks


def test_upstream_figure_parser_uses_injected_vision_model():
    from src.shared.utils.deepdoc.parser.figure_parser import vision_figure_parser_docx_wrapper

    image = Image.new("RGB", (32, 24), color="white")

    def fake_vision_model(binary, prompt):
        assert isinstance(binary, Image.Image)
        assert "Describe the figure" in prompt
        return "injected figure description"

    result = vision_figure_parser_docx_wrapper(
        [("caption", image)],
        [],
        vision_model=fake_vision_model,
    )

    assert len(result) == 1
    assert "injected figure description" in str(result[0][0][1])


def test_upstream_vision_parser_uses_injected_vision_model():
    from src.shared.utils.deepdoc.parser.pdf_parser import VisionParser

    def fake_vision_model(binary, prompt):
        assert "PDF page 1" in prompt
        return "vision parser text"

    parser = VisionParser(fake_vision_model)
    docs, tables = parser(_build_minimal_pdf_bytes("Vision Parser Source"), zoomin=1)

    assert tables == []
    assert docs
    assert docs[0][0] == "vision parser text"
    assert docs[0][1].startswith("@@1\t")


def test_deepdoc_parser_supports_parse_bytes_for_pdf():
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("Bytes PDF"),
            file_type="pdf",
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "pdf"
    assert "Bytes PDF" in result.full_text
    assert result.chunks


def test_deepdoc_parser_supports_plain_pdf_mode():
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("Plain PDF Mode"),
            file_type="pdf",
            parsing_config={"deepdoc_pdf_mode": "plain"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["file_type"] == "pdf"
    assert result.metadata["pdf_mode"] == "plain"
    assert "@@1\t" not in result.full_text


def test_deepdoc_parser_supports_opendataloader_pdf_parser(monkeypatch):
    class _FakeResponse:
        def raise_for_status(self):
            return None

        def json(self):
            return {
                "json_doc": {
                    "children": [
                        {"type": "title", "content": "OpenDataLoader Title"},
                        {"type": "paragraph", "content": "OpenDataLoader Body"},
                        {"type": "table", "html": "<table><tr><td>A</td></tr></table>"},
                    ]
                }
            }

    monkeypatch.setenv("OPENDATALOADER_APISERVER", "http://mock-opendataloader")
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.ragflow_opendataloader_parser.requests.post",
        lambda *args, **kwargs: _FakeResponse(),
    )

    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("OpenDataLoader PDF"),
            file_type="pdf",
            parsing_config={"deepdoc_parser_id": "pdf_opendataloader"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["parser_class"] == "RAGFlowOpenDataLoaderParser"
    assert "OpenDataLoader Title" in result.full_text
    assert result.metadata["tables"]
    assert "OpenDataLoader Body" in result.full_text


def test_deepdoc_parser_supports_docling_pdf_parser(monkeypatch):
    class _FakeResponse:
        status_code = 200

        def json(self):
            return [
                {"text": "Docling Chunk 1"},
                {"chunk": {"text": "Docling Chunk 2"}},
            ]

    monkeypatch.setenv("DOCLING_SERVER_URL", "http://mock-docling")
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.ragflow_docling_parser.requests.post",
        lambda *args, **kwargs: _FakeResponse(),
    )

    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("Docling PDF"),
            file_type="pdf",
            parsing_config={"deepdoc_parser_id": "pdf_docling"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["parser_class"] == "RAGFlowDoclingParser"
    assert result.metadata["docling_chunked"] is True
    assert "Docling Chunk 1" in result.full_text
    assert "Docling Chunk 2" in result.full_text


def test_deepdoc_parser_supports_mineru_pdf_parser(monkeypatch):
    class _FakeMinerUParser:
        def parse_bytes(self, file_bytes, *, parsing_config=None, file_name="input.pdf"):
            assert parsing_config["deepdoc_parser_id"] == "pdf_mineru"
            return (
                "@@1\t10.0\t120.0\t20.0\t40.0##MinerU Heading\n\n@@1\t10.0\t220.0\t50.0\t90.0##MinerU Body",
                [
                    "@@1\t10.0\t120.0\t20.0\t40.0##MinerU Heading",
                    "@@1\t10.0\t220.0\t50.0\t90.0##MinerU Body",
                ],
                {
                    "parser": "deepdoc",
                    "parser_class": "RAGFlowMinerUParser",
                    "file_type": "pdf",
                    "source": "ragflow-adapted",
                    "service": {"api_url": "http://mock-mineru", "configured": True},
                    "backend": "pipeline",
                    "tables": [],
                    "section_count": 2,
                },
            )

    monkeypatch.setenv("MINERU_APISERVER", "http://mock-mineru")

    parser = DeepDocParser()
    parser._mineru_parser = _FakeMinerUParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("MinerU PDF"),
            file_type="pdf",
            parsing_config={"deepdoc_parser_id": "pdf_mineru", "mineru_backend": "pipeline"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["parser_class"] == "RAGFlowMinerUParser"
    assert result.metadata["backend"] == "pipeline"
    assert "MinerU Heading" in result.full_text
    assert "MinerU Body" in result.full_text


def test_deepdoc_parser_supports_somark_pdf_parser(monkeypatch):
    class _FakeSoMarkParser:
        def parse_bytes(self, file_bytes, *, parsing_config=None, file_name="input.pdf"):
            assert parsing_config["deepdoc_parser_id"] == "pdf_somark"
            return (
                "@@1\t10.0\t180.0\t20.0\t60.0##SoMark Title\n\n@@1\t10.0\t220.0\t80.0\t120.0##SoMark Body",
                [
                    "@@1	10.0	180.0	20.0	60.0##SoMark Title",
                    "@@1	10.0	220.0	80.0	120.0##SoMark Body",
                ],
                {
                    "parser": "deepdoc",
                    "parser_class": "RAGFlowSoMarkParser",
                    "file_type": "pdf",
                    "source": "ragflow-adapted",
                    "service": {"base_url": "http://mock-somark", "configured": True},
                    "element_formats": {"image": "url", "formula": "latex", "table": "html", "cs": "image"},
                    "feature_config": {"enable_title_level_recognition": False},
                    "tables": [],
                    "section_count": 2,
                },
            )

    monkeypatch.setenv("SOMARK_BASE_URL", "http://mock-somark")

    parser = DeepDocParser()
    parser._somark_parser = _FakeSoMarkParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("SoMark PDF"),
            file_type="pdf",
            parsing_config={"deepdoc_parser_id": "pdf_somark"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["parser_class"] == "RAGFlowSoMarkParser"
    assert result.metadata["service"]["configured"] is True
    assert "SoMark Title" in result.full_text
    assert "SoMark Body" in result.full_text


def test_deepdoc_parser_supports_tcadp_pdf_parser(monkeypatch):
    class _FakeTCADPParser:
        def parse_bytes(self, file_bytes, *, parsing_config=None, file_name="input.pdf"):
            assert parsing_config["deepdoc_parser_id"] == "pdf_tcadp"
            return (
                "@@1\t0.0\t1000.0\t0.0\t100.0##TCADP Heading\n\n@@1\t0.0\t1000.0\t0.0\t100.0##TCADP Body",
                [
                    "@@1	0.0	1000.0	0.0	100.0##TCADP Heading",
                    "@@1	0.0	1000.0	0.0	100.0##TCADP Body",
                ],
                {
                    "parser": "deepdoc",
                    "parser_class": "RAGFlowTCADPParser",
                    "file_type": "pdf",
                    "source": "ragflow-adapted",
                    "service": {"region": "ap-guangzhou", "configured": True},
                    "table_result_type": "1",
                    "markdown_image_response_type": "1",
                    "tables": [],
                    "section_count": 2,
                },
            )

    monkeypatch.setenv("TCADP_SECRET_ID", "secret-id")
    monkeypatch.setenv("TCADP_SECRET_KEY", "secret-key")

    parser = DeepDocParser()
    parser._tcadp_parser = _FakeTCADPParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("TCADP PDF"),
            file_type="pdf",
            parsing_config={"deepdoc_parser_id": "pdf_tcadp"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["parser_class"] == "RAGFlowTCADPParser"
    assert result.metadata["service"]["configured"] is True
    assert "TCADP Heading" in result.full_text
    assert "TCADP Body" in result.full_text


def test_deepdoc_parser_supports_paddleocr_pdf_parser(monkeypatch):
    class _SubmitResponse:
        def raise_for_status(self):
            return None

        def json(self):
            return {"data": {"jobId": "job-123"}}

    class _PollResponse:
        def raise_for_status(self):
            return None

        def json(self):
            return {"data": {"state": "done", "resultJsonUrl": "http://mock-paddleocr/result.jsonl"}}

    class _ResultResponse:
        headers = {"content-type": "application/jsonl"}
        text = '{"result":{"layoutParsingResults":[{"prunedResult":{"parsing_res_list":[{"block_content":"# PaddleOCR Title","block_bbox":[10,20,210,80]},{"block_content":"PaddleOCR Body","block_bbox":[12,100,220,160]}]}}],"ocrResults":[]}}'

        def raise_for_status(self):
            return None

        def json(self):
            raise AssertionError("json() should not be used for jsonl payloads")

    monkeypatch.setenv("PADDLEOCR_BASE_URL", "http://mock-paddleocr")

    def _fake_post(url, *args, **kwargs):
        assert url.endswith("/api/v2/ocr/jobs")
        assert kwargs["data"]["model"]
        return _SubmitResponse()

    def _fake_get(url, *args, **kwargs):
        if url.endswith("/api/v2/ocr/jobs/job-123"):
            return _PollResponse()
        if url == "http://mock-paddleocr/result.jsonl":
            return _ResultResponse()
        raise AssertionError(f"unexpected url: {url}")

    monkeypatch.setattr(
        "src.shared.utils.deepdoc.ragflow_paddleocr_parser.requests.post",
        _fake_post,
    )
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.ragflow_paddleocr_parser.requests.get",
        _fake_get,
    )

    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("PaddleOCR PDF"),
            file_type="pdf",
            parsing_config={"deepdoc_parser_id": "pdf_paddleocr"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["parser_class"] == "RAGFlowPaddleOCRParser"
    assert result.metadata["algorithm"]
    assert "PaddleOCR Title" in result.full_text
    assert "PaddleOCR Body" in result.full_text
    assert "@@1	" in result.full_text


def test_ragflow_pdf_parser_parse_into_bboxes():
    parser = RAGFlowPdfParser()
    result = parser(_build_minimal_pdf_bytes("BBox PDF"), pdf_mode="layout", chunk_size=500)
    assert result.metadata["parser_class"] == "RAGFlowPdfParser"
    assert result.metadata["bboxes"]
    assert "BBox PDF" in result.full_text


def test_ragflow_pdf_parser_detects_columns_and_positions():
    parser = RAGFlowPdfParser()
    pdf_bytes = _build_positioned_pdf_bytes(
        [
            (72, 700, "Left One"),
            (72, 660, "Left Two"),
            (320, 700, "Right One"),
            (320, 660, "Right Two"),
        ]
    )
    result = parser(pdf_bytes, pdf_mode="layout", chunk_size=500)

    assert result.metadata["detected_columns"] == 2
    assert result.full_text.index("Left One") < result.full_text.index("Right One")
    assert result.full_text.index("Left Two") < result.full_text.index("Right One")
    assert result.metadata["bboxes"][0]["position_tag"].startswith("@@1\t")
    assert result.metadata["merged_block_count"] <= len(result.metadata["bboxes"])


def test_ragflow_pdf_parser_merges_vertical_lines_into_blocks():
    parser = RAGFlowPdfParser()
    pdf_bytes = _build_positioned_pdf_bytes(
        [
            (72, 700, "This is line one"),
            (72, 684, "that continues below"),
            (72, 630, "New paragraph"),
        ]
    )
    result = parser(pdf_bytes, pdf_mode="layout", chunk_size=500)

    assert "This is line one that continues below" in result.full_text
    assert result.metadata["merged_block_count"] < len(result.metadata["bboxes"])
    assert result.metadata["merged_bboxes"][0]["positions"]
    assert result.metadata["paragraph_merge_strategy"] in {"heuristic", "xgboost"}
    assert "page_filter" in result.metadata
    assert "artifacts" in result.metadata
    assert "table_regions" in result.metadata
    assert "reading_order" in result.metadata
    assert "chunk_structure" in result.metadata


def test_structured_chunk_builder_respects_region_boundaries():
    reading_order = [
        {"kind": "text", "page": 1, "bbox": {"top": 10.0, "x0": 10.0}, "text": "Alpha", "source_id": "t1"},
        {"kind": "table", "page": 1, "bbox": {"top": 20.0, "x0": 10.0}, "caption": "Table 1", "text": "Row A", "source_id": "tb1"},
        {"kind": "figure", "page": 1, "bbox": {"top": 30.0, "x0": 10.0}, "caption": "Figure 1", "text": "Chart", "source_id": "fg1"},
    ]

    chunks, chunk_structure = RAGFlowPdfParser._build_structured_chunks(reading_order, chunk_size=20)

    assert len(chunks) == 3
    assert chunks[0] == "Alpha"
    assert chunks[1].startswith("[TABLE]")
    assert chunks[2].startswith("[FIGURE]")
    assert chunk_structure[0]["entry_kinds"] == ["text"]
    assert chunk_structure[1]["entry_kinds"] == ["table"]
    assert chunk_structure[2]["entry_kinds"] == ["figure"]


def test_text_concat_model_status_reports_missing_model_by_default():
    status = get_text_concat_model_status()

    assert status["filename"] == "updown_concat_xgb.model"
    assert "repo_id" in status
    assert "available" in status


def test_updown_concat_merger_falls_back_without_model():
    merger = UpDownConcatMerger()
    boxes = [
        RAGFlowPdfParser().parse_into_bboxes(_build_positioned_pdf_bytes([(72, 700, "Line One"), (72, 684, "Line Two")]))
    ][0]

    merged, strategy = merger.merge(boxes)

    assert strategy == "heuristic"
    assert len(merged) == 1
    assert "Line One Line Two" in merged[0].text


def test_updown_concat_merger_can_use_mocked_xgb_model(monkeypatch):
    if not get_deepdoc_runtime_report()["xgboost"]["available"]:
        pytest.skip("xgboost is not available in the current runtime")

    merger = UpDownConcatMerger()
    boxes = RAGFlowPdfParser().parse_into_bboxes(_build_positioned_pdf_bytes([(72, 700, "Line One"), (72, 684, "Line Two")]))

    class _FakeModel:
        def predict(self, matrix):
            return np.array([1.0], dtype=np.float32)

    monkeypatch.setattr(merger, "model_available", lambda: True)
    monkeypatch.setattr(merger, "load_model", lambda: _FakeModel())

    merged, strategy = merger.merge(boxes)

    assert strategy == "xgboost"
    assert len(merged) == 1
    assert "Line One Line Two" in merged[0].text


def test_page_noise_filter_removes_toc_section():
    box_cls = DeepDocPdfBox
    boxes = [
        box_cls(page=1, x0=10, x1=50, top=10, bottom=20, text="Contents"),
        box_cls(page=1, x0=10, x1=60, top=30, bottom=40, text="1 Intro"),
        box_cls(page=1, x0=10, x1=80, top=50, bottom=60, text="2 Method"),
        box_cls(page=2, x0=10, x1=70, top=10, bottom=20, text="Intro"),
        box_cls(page=2, x0=10, x1=120, top=30, bottom=40, text="Body starts here"),
    ]

    filtered, meta = PageNoiseFilter().filter_boxes(boxes, total_pages=2)

    assert meta["toc_detected"] is True
    assert meta["removed_boxes"] >= 2
    assert all("Contents" not in box.text for box in filtered)
    assert any(box.page == 2 for box in filtered)


def test_page_noise_filter_removes_dirty_pages():
    box_cls = DeepDocPdfBox
    boxes = [
        box_cls(page=1, x0=10, x1=60, top=10, bottom=20, text="(cid:12)"),
        box_cls(page=1, x0=10, x1=60, top=30, bottom=40, text="(cid:13)"),
        box_cls(page=1, x0=10, x1=60, top=50, bottom=60, text="(cid:14)"),
        box_cls(page=1, x0=10, x1=60, top=70, bottom=80, text="(cid:15)"),
        box_cls(page=2, x0=10, x1=80, top=10, bottom=20, text="Clean page"),
    ]

    filtered, meta = PageNoiseFilter().filter_boxes(boxes, total_pages=2)

    assert meta["toc_detected"] is False
    assert meta["dirty_pages"] == [1]
    assert all(box.page != 1 for box in filtered)
    assert filtered[0].text == "Clean page"


def test_ragflow_pdf_parser_applies_page_filter(monkeypatch):
    parser = RAGFlowPdfParser()

    monkeypatch.setattr(
        parser,
        "parse_into_bboxes",
        lambda filename: [
            DeepDocPdfBox(page=1, x0=10, x1=50, top=10, bottom=20, text="Contents"),
            DeepDocPdfBox(page=1, x0=10, x1=60, top=30, bottom=40, text="1 Intro"),
            DeepDocPdfBox(page=2, x0=10, x1=70, top=10, bottom=20, text="Actual body"),
        ],
    )
    monkeypatch.setattr(
        parser,
        "_merge_vertical_boxes_with_strategy",
        lambda boxes: (list(boxes), "heuristic"),
    )

    result = parser(_build_minimal_pdf_bytes("Body"), pdf_mode="layout", chunk_size=500)

    assert result.metadata["page_filter"]["toc_detected"] is True
    assert "Contents" not in result.full_text
    assert "Actual body" in result.full_text


def test_pdf_artifact_extractor_groups_tables_and_figures():
    _skip_if_vision_runtime_unavailable()
    extractor = PdfArtifactExtractor()
    page_image = Image.new("RGB", (400, 200), color=(255, 255, 255))
    boxes = [
        DeepDocPdfBox(page=1, x0=20, x1=180, top=20, bottom=35, text="Table 1 Results", layout_type="table caption"),
        DeepDocPdfBox(page=1, x0=20, x1=100, top=40, bottom=55, text="Metric | Value", layout_type="table"),
        DeepDocPdfBox(page=1, x0=20, x1=100, top=58, bottom=72, text="Recall | 0.88", layout_type="table"),
        DeepDocPdfBox(page=1, x0=220, x1=360, top=20, bottom=35, text="Figure 1 Trend", layout_type="figure caption"),
        DeepDocPdfBox(page=1, x0=220, x1=360, top=40, bottom=140, text="Chart region", layout_type="figure"),
    ]

    artifacts = extractor.extract(boxes, page_images={1: page_image}, zoom=1.0)

    assert len(artifacts["tables"]) == 1
    assert len(artifacts["figures"]) == 1
    assert artifacts["tables"][0]["caption"] == "Table 1 Results"
    assert "<table>" in artifacts["tables"][0]["html"]
    assert "<th>Metric</th>" in artifacts["tables"][0]["html"]
    assert artifacts["tables"][0]["html_source"] == "tsr_constructed"
    assert artifacts["tables"][0]["has_image"] is True
    assert isinstance(artifacts["tables"][0]["image"], LazyImage)
    assert artifacts["tables"][0]["image"].first()
    assert artifacts["figures"][0]["caption"] == "Figure 1 Trend"
    assert artifacts["figures"][0]["has_image"] is True

    figure_regions = RAGFlowPdfParser._build_figure_regions_metadata(artifacts)
    assert len(figure_regions) == 1
    assert figure_regions[0]["page_start"] == 1
    assert figure_regions[0]["region_index_on_page"] == 0
    assert figure_regions[0]["member_text_count"] == 2


def test_pdf_artifact_extractor_stitches_cross_page_crops():
    _skip_if_vision_runtime_unavailable()
    extractor = PdfArtifactExtractor()
    page1 = Image.new("RGB", (240, 200), color=(255, 255, 255))
    page2 = Image.new("RGB", (240, 200), color=(240, 240, 240))
    boxes = [
        DeepDocPdfBox(page=1, x0=20, x1=180, top=40, bottom=90, text="Part A", layout_type="figure"),
        DeepDocPdfBox(page=2, x0=20, x1=180, top=20, bottom=70, text="Part B", layout_type="figure"),
    ]

    artifacts = extractor.extract(boxes, page_images={1: page1, 2: page2}, zoom=1.0)

    assert len(artifacts["figures"]) == 1
    image = artifacts["figures"][0]["image"]
    assert isinstance(image, LazyImage)
    assert len(image.blobs) == 3
    stitched = Image.open(BytesIO(image.blobs[0]))
    first_page_crop = Image.open(BytesIO(image.blobs[1]))
    second_page_crop = Image.open(BytesIO(image.blobs[2]))
    assert stitched.size[1] == first_page_crop.size[1] + second_page_crop.size[1]
    assert stitched.size[0] == max(first_page_crop.size[0], second_page_crop.size[0])


def test_pdf_artifact_extractor_constructs_table_html_from_aligned_boxes():
    _skip_if_vision_runtime_unavailable()
    extractor = PdfArtifactExtractor()
    boxes = [
        DeepDocPdfBox(page=1, x0=20, x1=80, top=40, bottom=55, text="Metric", layout_type="table"),
        DeepDocPdfBox(page=1, x0=90, x1=150, top=40, bottom=55, text="Value", layout_type="table"),
        DeepDocPdfBox(page=1, x0=20, x1=80, top=60, bottom=75, text="Recall", layout_type="table"),
        DeepDocPdfBox(page=1, x0=90, x1=150, top=60, bottom=75, text="0.88", layout_type="table"),
    ]

    artifacts = extractor.extract(boxes)

    assert len(artifacts["tables"]) == 1
    assert artifacts["tables"][0]["html_source"] == "tsr_constructed"
    assert "<th>Metric</th>" in artifacts["tables"][0]["html"]
    assert "<td>Recall</td>" in artifacts["tables"][0]["html"]
    assert "<td>0.88</td>" in artifacts["tables"][0]["html"]


def test_pdf_artifact_extractor_can_use_mocked_tsr_model():
    _skip_if_vision_runtime_unavailable()
    extractor = PdfArtifactExtractor()
    page_image = Image.new("RGB", (220, 160), color=(255, 255, 255))
    boxes = [
        DeepDocPdfBox(page=1, x0=20, x1=80, top=40, bottom=55, text="Metric", layout_type="table"),
        DeepDocPdfBox(page=1, x0=90, x1=150, top=40, bottom=55, text="Value", layout_type="table"),
        DeepDocPdfBox(page=1, x0=20, x1=80, top=60, bottom=75, text="Recall", layout_type="table"),
        DeepDocPdfBox(page=1, x0=90, x1=150, top=60, bottom=75, text="0.88", layout_type="table"),
    ]

    class _FakeTsr:
        def __call__(self, images, thr=0.2):
            assert len(images) == 1
            return [[
                {"label": "table row", "score": 0.95, "x0": 0.0, "x1": 130.0, "top": 0.0, "bottom": 18.0},
                {"label": "table row", "score": 0.94, "x0": 0.0, "x1": 130.0, "top": 18.0, "bottom": 35.0},
                {"label": "table column", "score": 0.95, "x0": 0.0, "x1": 65.0, "top": 0.0, "bottom": 35.0},
                {"label": "table column", "score": 0.95, "x0": 65.0, "x1": 130.0, "top": 0.0, "bottom": 35.0},
                {"label": "table column header", "score": 0.9, "x0": 0.0, "x1": 130.0, "top": 0.0, "bottom": 18.0},
            ]]

    extractor._tsr_attempted = True
    extractor._tsr = _FakeTsr()

    artifacts = extractor.extract(boxes, page_images={1: page_image}, zoom=1.0)

    assert artifacts["tables"][0]["html_source"] == "tsr_model"
    assert artifacts["tables"][0]["table_structure"]["source"] == "tsr_model"
    assert artifacts["tables"][0]["table_structure"]["prediction_pages"] == 1
    assert artifacts["tables"][0]["table_structure"]["prediction_count"] == 5
    assert len(artifacts["tables"][0]["table_structure"]["structured_boxes"]) == 4
    assert "<th>Metric</th>" in artifacts["tables"][0]["html"]
    assert "<td>Recall</td>" in artifacts["tables"][0]["html"]

    regions = RAGFlowPdfParser._build_table_regions_metadata(artifacts)
    assert len(regions) == 1
    assert regions[0]["page_start"] == 1
    assert regions[0]["region_index_on_page"] == 0
    assert regions[0]["member_text_count"] == 4
    assert regions[0]["row_count"] == 2
    assert regions[0]["column_count"] == 2
    assert regions[0]["structured_box_count"] == 4
    assert regions[0]["table_structure_source"] == "tsr_model"

    reading_order = RAGFlowPdfParser._build_reading_order_metadata(
        [
            DeepDocPdfBox(page=1, x0=5, x1=15, top=5, bottom=10, text="Intro", position_tag="@@1\t5\t15\t5\t10##"),
        ],
        regions,
        [],
    )
    assert [item["kind"] for item in reading_order] == ["text", "table"]
    assert reading_order[0]["global_order"] == 0
    assert reading_order[1]["global_order"] == 1
    assert reading_order[0]["order_on_page"] == 0
    assert reading_order[1]["order_on_page"] == 1


def test_deepdoc_parser_vision_mode_surfaces_artifacts():
    _skip_if_vision_runtime_unavailable()
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("Vision Mode"),
            file_type="pdf",
            parsing_config={"deepdoc_pdf_mode": "vision"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert "artifacts" in result.metadata
    assert "tables" in result.metadata["artifacts"]
    assert "figures" in result.metadata["artifacts"]


def test_ragflow_pdf_parser_layout_artifacts_include_lazy_images():
    parser = RAGFlowPdfParser()
    result = parser(_build_minimal_pdf_bytes("Layout Artifact"), pdf_mode="layout", chunk_size=500)

    assert "artifacts" in result.metadata
    assert isinstance(result.metadata["artifacts"]["tables"], list)
    assert isinstance(result.metadata["artifacts"]["figures"], list)


def test_ragflow_pdf_parser_position_helpers():
    text = "@@1\t10.0\t20.0\t30.0\t40.0##hello\n@@2\t11.0\t21.0\t31.0\t41.0##world"
    assert RAGFlowPdfParser.remove_tag(text) == "hello\nworld"

    positions = RAGFlowPdfParser.extract_positions(text)
    assert positions[0][0] == [0]
    assert positions[0][1:] == (10.0, 20.0, 30.0, 40.0)
    assert positions[1][0] == [1]


def test_deepdoc_engine_standalone_facade():
    engine = DeepDocEngine()
    assert engine.can_parse("pdf") is True
    assert engine.can_parse(".docx") is True
    assert engine.can_parse("epub") is True
    assert engine.can_parse("xlsx") is True
    assert engine.can_parse("pptx") is True
    assert engine.can_parse("png") is True
    assert engine.supports_pdf_mode("plain") is True
    assert engine.supports_pdf_mode("vision") is get_vision_runtime_status()["parser_available"]


def test_deepdoc_factory_exposes_parser_ids():
    specs = DeepDocParserFactory.list_specs()
    assert "pdf_layout" in specs
    assert "pdf_plain" in specs
    assert "pdf_vision" in specs
    assert "pdf_docling" in specs
    assert "pdf_mineru" in specs
    assert "pdf_opendataloader" in specs
    assert "pdf_paddleocr" in specs
    assert "pdf_somark" in specs
    assert "pdf_tcadp" in specs
    assert "epub" in specs
    assert "excel" in specs
    assert "ppt" in specs
    assert "figure" in specs
    assert "markdown" in specs
    assert "html" in specs
    assert "json" in specs
    assert specs["pdf_plain"].mode == "plain"
    assert specs["pdf_vision"].available is get_vision_runtime_status()["parser_available"]


def test_knowledge_base_config_accepts_deepdoc_parser_id():
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_parser_id": "pdf_layout"},
        }
    )

    assert config.model_dump()["parsing"]["deepdoc_parser_id"] == "pdf_layout"


def test_knowledge_base_config_accepts_mineru_parser_id():
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_parser_id": "pdf_mineru"},
        }
    )

    assert config.model_dump()["parsing"]["deepdoc_parser_id"] == "pdf_mineru"


def test_knowledge_base_config_accepts_opendataloader_parser_id():
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_parser_id": "pdf_opendataloader"},
        }
    )

    assert config.model_dump()["parsing"]["deepdoc_parser_id"] == "pdf_opendataloader"


def test_knowledge_base_config_accepts_docling_parser_id():
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_parser_id": "pdf_docling"},
        }
    )

    assert config.model_dump()["parsing"]["deepdoc_parser_id"] == "pdf_docling"


def test_knowledge_base_config_accepts_tcadp_parser_id():
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_parser_id": "pdf_tcadp"},
        }
    )

    assert config.model_dump()["parsing"]["deepdoc_parser_id"] == "pdf_tcadp"


def test_knowledge_base_config_accepts_somark_parser_id():
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_parser_id": "pdf_somark"},
        }
    )

    assert config.model_dump()["parsing"]["deepdoc_parser_id"] == "pdf_somark"


def test_knowledge_base_config_accepts_paddleocr_parser_id():
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_parser_id": "pdf_paddleocr"},
        }
    )

    assert config.model_dump()["parsing"]["deepdoc_parser_id"] == "pdf_paddleocr"


@pytest.mark.parametrize(
    "parser_id",
    ["epub", "excel", "ppt", "figure", "text", "txt", "markdown", "html", "json"],
)
def test_knowledge_base_config_accepts_non_pdf_deepdoc_parser_ids(parser_id: str):
    config = KnowledgeBaseConfig.model_validate(
        {
            "parsing": {"strategy": "deepdoc", "deepdoc_parser_id": parser_id},
        }
    )

    assert config.model_dump()["parsing"]["deepdoc_parser_id"] == parser_id


@pytest.mark.parametrize(
    "parser_id",
    ["epub", "excel", "ppt", "figure", "text", "txt", "markdown", "html", "json"],
)
def test_knowledge_base_service_accepts_non_pdf_deepdoc_parser_ids(parser_id: str):
    service = object.__new__(KnowledgeBaseService)

    service._validate_config_updates(
        {
            "parsing": {
                "strategy": "deepdoc",
                "deepdoc_parser_id": parser_id,
            }
        }
    )


def test_deepdoc_capabilities_surface_unavailable_vision_mode():
    capabilities = get_deepdoc_capabilities()
    assert "pdf_modes" in capabilities
    assert capabilities["pdf_modes"]["plain"]["available"] is True
    assert capabilities["pdf_modes"]["layout"]["available"] is True
    assert capabilities["pdf_modes"]["vision"]["available"] is get_vision_runtime_status()["parser_available"]
    assert "vision_runtime" in capabilities
    assert "package_status" in capabilities["pdf_modes"]["vision"]


def test_deepdoc_runtime_report_surfaces_missing_heavy_dependencies():
    report = get_deepdoc_runtime_report()
    assert report["pdfplumber"]["available"] is True
    assert "xgboost" in report
    assert report["pypdf"]["available"] is True
    assert "cv2" in report
    assert report["pandas"]["available"] is True
    assert "openpyxl" in report
    assert "pptx" in report


def test_vision_runtime_status_tracks_required_and_optional_dependencies():
    status = get_vision_runtime_status()
    assert status["available"] is (len(status["missing_required"]) == 0)
    assert status["parser_available"] is (status["available"] and status["package_status"]["implementation_ready"])
    assert "model_status" in status
    assert "required_dependencies" in status
    assert "optional_dependencies" in status


def test_vision_health_status_surfaces_model_groups():
    status = get_vision_health_status()
    assert "model_groups" in status
    assert "ocr" in status["model_groups"]
    assert "can_run_pdf_vision" in status


def test_vision_smoke_check_surfaces_check_matrix():
    status = run_vision_smoke_check()
    assert "ok" in status
    assert "checks" in status
    assert "pdf_vision_parser" in status["checks"]
    assert "load_checks" in status
    assert "inference_checks" in status
    assert "layout_model_load" in status["load_checks"]
    assert "layout_model_inference" in status["inference_checks"]


def test_vision_smoke_check_attempts_component_loads(monkeypatch):
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime.get_vision_health_status",
        lambda: {
            "runtime_available": True,
            "parser_available": True,
            "model_dir": "C:/models",
            "model_repo_id": "InfiniFlow/deepdoc",
            "required_missing": [],
            "optional_missing": [],
            "model_groups": {},
            "can_run_pdf_vision": True,
            "can_run_vendored_ocr": True,
            "can_run_layout_inference": True,
            "can_run_tsr_inference": True,
        },
    )
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime._attempt_component_load",
        lambda component: {"attempted": True, "ok": True, "error": None, "component": component},
    )
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime._attempt_component_inference",
        lambda component: {"attempted": True, "ok": True, "error": None, "component": component},
    )
    status = run_vision_smoke_check()
    assert status["load_checks"]["vendored_ocr_load"]["ok"] is True
    assert status["load_checks"]["layout_model_load"]["ok"] is True
    assert status["load_checks"]["tsr_model_load"]["ok"] is True
    assert status["inference_checks"]["vendored_ocr_inference"]["ok"] is True
    assert status["inference_checks"]["layout_model_inference"]["ok"] is True
    assert status["inference_checks"]["tsr_model_inference"]["ok"] is True


def test_vision_smoke_check_skips_component_loads_when_models_unavailable(monkeypatch):
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime.get_vision_health_status",
        lambda: {
            "runtime_available": True,
            "parser_available": True,
            "model_dir": "C:/models",
            "model_repo_id": "InfiniFlow/deepdoc",
            "required_missing": [],
            "optional_missing": [],
            "model_groups": {},
            "can_run_pdf_vision": True,
            "can_run_vendored_ocr": False,
            "can_run_layout_inference": False,
            "can_run_tsr_inference": False,
        },
    )
    status = run_vision_smoke_check()
    assert status["load_checks"]["vendored_ocr_load"]["attempted"] is False
    assert status["load_checks"]["layout_model_load"]["attempted"] is False
    assert status["load_checks"]["tsr_model_load"]["attempted"] is False
    assert status["inference_checks"]["vendored_ocr_inference"]["attempted"] is False
    assert status["inference_checks"]["layout_model_inference"]["attempted"] is False
    assert status["inference_checks"]["tsr_model_inference"]["attempted"] is False


def test_vision_smoke_check_skips_inference_when_load_fails(monkeypatch):
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime.get_vision_health_status",
        lambda: {
            "runtime_available": True,
            "parser_available": True,
            "model_dir": "C:/models",
            "model_repo_id": "InfiniFlow/deepdoc",
            "required_missing": [],
            "optional_missing": [],
            "model_groups": {},
            "can_run_pdf_vision": True,
            "can_run_vendored_ocr": True,
            "can_run_layout_inference": True,
            "can_run_tsr_inference": True,
        },
    )
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime._attempt_component_load",
        lambda component: {"attempted": True, "ok": component == "layout", "error": None},
    )
    called = []

    def _fake_inference(component):
        called.append(component)
        return {"attempted": True, "ok": True, "error": None}

    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime._attempt_component_inference",
        _fake_inference,
    )

    status = run_vision_smoke_check()

    assert called == ["layout"]
    assert status["inference_checks"]["vendored_ocr_inference"]["attempted"] is False
    assert status["inference_checks"]["layout_model_inference"]["ok"] is True
    assert status["inference_checks"]["tsr_model_inference"]["attempted"] is False


def test_attempt_component_inference_supports_mocked_layout_and_tsr(monkeypatch):
    from src.shared.utils.deepdoc import vision_runtime

    class _FakeLayoutRecognizer:
        def __init__(self, *, autoload):
            assert autoload is True

        def forward(self, images, thr=0.0, batch_size=1):
            assert len(images) == 1
            return [[{"type": "Title", "score": 0.9, "bbox": [0, 0, 10, 10]}]]

    class _FakeTsrRecognizer:
        def __init__(self, *, autoload):
            assert autoload is True

        def forward(self, images, thr=0.0, batch_size=1):
            assert len(images) == 1
            return [[{"type": "table row", "score": 0.9, "bbox": [0, 0, 10, 10]}]]

    monkeypatch.setitem(sys.modules, "src.shared.utils.deepdoc.vision.layout_recognizer", SimpleNamespace(LayoutRecognizer=_FakeLayoutRecognizer))
    monkeypatch.setitem(
        sys.modules,
        "src.shared.utils.deepdoc.vision.table_structure_recognizer",
        SimpleNamespace(TableStructureRecognizer=_FakeTsrRecognizer),
    )

    layout_result = vision_runtime._attempt_component_inference("layout")
    tsr_result = vision_runtime._attempt_component_inference("tsr")

    assert layout_result["ok"] is True
    assert layout_result["pages"] == 1
    assert tsr_result["ok"] is True
    assert tsr_result["pages"] == 1


def test_vision_runtime_guard_raises_clear_error():
    status = get_vision_runtime_status()
    if status["available"]:
        assert ensure_vision_runtime_available()["available"] is True
        parser_status = ensure_vision_parser_available()
        assert parser_status["parser_available"] is True
    else:
        with pytest.raises(DeepDocVisionRuntimeUnavailable):
            ensure_vision_runtime_available()
        with pytest.raises(DeepDocVisionRuntimeUnavailable):
            ensure_vision_parser_available()


def test_resume_surname_compatibility_preserves_chinese_names():
    from src.shared.integrations.deepdoc.compat.compat import surname

    assert surname.isit("\u738b") is True
    assert surname.isit("\u6b27\u9633") is True
    assert surname.isit("Alice") is False


def test_vendored_vision_seeit_draws_and_saves_results(tmp_path):
    from src.shared.utils.deepdoc.vision.seeit import draw_box, save_results

    image = Image.new("RGB", (80, 60), color="white")
    detections = [{"type": "title", "score": 0.95, "bbox": [5, 8, 50, 30]}]

    rendered = draw_box(image, detections, labels=["title"], threshold=0.5)
    saved = save_results([image], [detections], labels=["title"], output_dir=tmp_path)

    assert rendered.size == image.size
    assert rendered.getpixel((5, 8)) != (255, 255, 255)
    assert saved == [tmp_path / "0.jpg"]
    assert saved[0].exists()


def test_upstream_snapshot_matches_implemented_server_and_vision_modules():
    snapshot = get_upstream_deepdoc_snapshot()

    assert snapshot["server_modules"]["missing"] == []
    assert "docker_stubs" in snapshot["server_modules"]["implemented"]
    assert "seeit" in snapshot["vision_modules"]["implemented"]
    assert "seeit" not in snapshot["vision_modules"]["missing"]
    assert snapshot["vision_modules"]["missing"] == []


def test_vendored_docker_stubs_write_minimal_packages(tmp_path):
    from src.shared.utils.deepdoc.server.docker_stubs import write_docker_stubs

    written = write_docker_stubs(tmp_path)

    assert tmp_path / "deepdoc" / "__init__.py" in written
    assert tmp_path / "deepdoc" / "vision" / "__init__.py" in written
    assert tmp_path / "common" / "file_utils.py" in written
    assert tmp_path / "rag" / "nlp" / "__init__.py" in written
    assert tmp_path / "rag" / "utils" / "lazy_image.py" in written
    assert all(path.exists() for path in written)


def test_vendored_ocr_diagnostic_entrypoint(tmp_path):
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.t_ocr import run_ocr_diagnostics

    image_path = tmp_path / "ocr.png"
    Image.new("RGB", (64, 48), color="white").save(image_path)

    class _FakeOCR:
        def __call__(self, image):
            return [
                (
                    [[5, 6], [35, 6], [35, 20], [5, 20]],
                    ("hello", 0.95),
                )
            ]

    outputs = run_ocr_diagnostics(image_path, tmp_path / "ocr-out", ocr=_FakeOCR())

    assert len(outputs) == 1
    assert outputs[0]["detections"][0]["text"] == "hello"
    assert outputs[0]["image"].exists()
    assert outputs[0]["text"].read_text(encoding="utf-8") == "hello"


def test_vendored_recognizer_diagnostic_entrypoints(tmp_path):
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.t_recognizer import run_recognizer_diagnostics

    image_path = tmp_path / "layout.png"
    Image.new("RGB", (64, 48), color="white").save(image_path)

    class _FakeLayout:
        labels = ["title"]

        def forward(self, images, thr=0.5, batch_size=16):
            return [[{"type": "title", "score": 0.9, "bbox": [4, 5, 40, 22]}]]

    class _FakeTSR:
        labels = ["table row"]

        def __call__(self, images, thr=0.5):
            return [[{"label": "table row", "score": 0.88, "x0": 3, "top": 8, "x1": 50, "bottom": 24}]]

    layout_outputs = run_recognizer_diagnostics(
        image_path,
        tmp_path / "layout-out",
        mode="layout",
        recognizer=_FakeLayout(),
    )
    tsr_outputs = run_recognizer_diagnostics(
        image_path,
        tmp_path / "tsr-out",
        mode="tsr",
        recognizer=_FakeTSR(),
    )

    assert layout_outputs[0]["detections"][0]["type"] == "title"
    assert layout_outputs[0]["image"].exists()
    assert layout_outputs[0]["json"].exists()
    assert tsr_outputs[0]["detections"][0]["type"] == "table row"
    assert tsr_outputs[0]["image"].exists()
    assert tsr_outputs[0]["json"].exists()


def test_upstream_snapshot_has_no_directory_level_gaps():
    snapshot = get_upstream_deepdoc_snapshot()

    assert snapshot["parser_modules"]["missing"] == []
    assert snapshot["vision_modules"]["missing"] == []
    assert snapshot["server_modules"]["missing"] == []


def test_vendored_vision_package_status_exposed():
    status = get_vendored_vision_package_status()
    assert status["present"] is True
    assert status["implementation_ready"] is True
    assert "model_status" in status
    assert "recognizer.py" in status["modules"]
    assert "ocr.py" in status["modules"]
    assert "layout_recognizer.py" in status["modules"]
    assert "table_structure_recognizer.py" in status["modules"]
    assert "operators.py" in status["modules"]
    assert "postprocess.py" in status["modules"]


def test_vendored_vision_recognizer_geometry_helpers():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision import Recognizer as DeepDocVisionRecognizer

    boxes = [
        {"x0": 50, "x1": 100, "top": 80, "bottom": 100},
        {"x0": 10, "x1": 60, "top": 10, "bottom": 30},
        {"x0": 12, "x1": 62, "top": 40, "bottom": 60},
    ]
    sorted_boxes = DeepDocVisionRecognizer.sort_Y_firstly(boxes, threshold=5)
    assert sorted_boxes[0]["top"] == 10
    assert sorted_boxes[1]["top"] == 40
    assert DeepDocVisionRecognizer.overlapped_area(
        {"x0": 0, "x1": 10, "top": 0, "bottom": 10},
        {"x0": 5, "x1": 15, "top": 0, "bottom": 10},
    ) > 0


def test_vendored_vision_recognizer_load_raises_for_missing_model():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision import Recognizer as DeepDocVisionRecognizer

    recognizer = DeepDocVisionRecognizer(labels=["Text"], domain="layout", model_dir=Path("C:/missing-model-dir"))
    with pytest.raises(FileNotFoundError):
        recognizer.load()


def test_vendored_vision_model_manager_reports_expected_files(tmp_path):
    (tmp_path / "det.onnx").write_bytes(b"det")
    status = get_model_status(tmp_path)

    assert "layout.onnx" in expected_model_files("layout")
    assert "det.onnx" in expected_model_files("ocr")
    assert status["groups"]["ocr"]["available"] is False
    assert "rec.onnx" in status["groups"]["ocr"]["missing"]
    assert status["groups"]["layout"]["available"] is False


def test_vendored_vision_model_manager_raises_for_missing_group(tmp_path):
    with pytest.raises(FileNotFoundError):
        ensure_model_group_available("ocr", tmp_path)


def test_vendored_vision_operator_helpers():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.operators import KeepKeys, create_operators, nms

    operators = create_operators([{"KeepKeys": {"keep_keys": ["a", "b"]}}])
    assert len(operators) == 1
    assert isinstance(operators[0], KeepKeys)
    assert operators[0]({"a": 1, "b": 2, "c": 3}) == [1, 2]

    kept = nms(
        np.array([[0, 0, 10, 10], [1, 1, 9, 9], [20, 20, 30, 30]], dtype=np.float32),
        np.array([0.9, 0.8, 0.7], dtype=np.float32),
        0.3,
    )
    assert kept[0] == 0
    assert 2 in kept


def test_vendored_vision_postprocess_builders():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.postprocess import build_post_process

    decoder = build_post_process({"name": "CTCLabelDecode"})
    assert decoder is not None
    post = build_post_process({"name": "DBPostProcess"})
    assert post is not None


def test_vendored_vision_ocr_helpers_without_model_load():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.ocr import OCR as DeepDocVisionOCR

    ocr = DeepDocVisionOCR(autoload=False)
    assert ".cache" in str(ocr.model_dir)

    boxes = np.array(
        [
            [[40, 40], [60, 40], [60, 60], [40, 60]],
            [[10, 10], [30, 10], [30, 30], [10, 30]],
        ],
        dtype=np.float32,
    )
    sorted_boxes = ocr.sorted_boxes(boxes)
    assert sorted_boxes[0][0][0] == 10

    img = np.zeros((80, 80, 3), dtype=np.uint8)
    pts = np.array([[10, 10], [30, 10], [30, 30], [10, 30]], dtype=np.float32)
    cropped = ocr.get_rotate_crop_image(img, pts)
    assert cropped.shape[0] > 0
    assert cropped.shape[1] > 0


def test_vendored_layout_recognizer_can_apply_layouts():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.layout_recognizer import LayoutRecognizer

    recognizer = LayoutRecognizer()
    image = np.zeros((300, 300, 3), dtype=np.uint8)
    ocr_res = [[
        {"text": "Report Title", "x0": 10, "x1": 120, "top": 10, "bottom": 25},
        {"text": "Body paragraph", "x0": 10, "x1": 150, "top": 80, "bottom": 100},
    ]]
    layouts = [[
        {"type": "title", "score": 0.9, "bbox": [0, 0, 360, 90]},
        {"type": "text", "score": 0.95, "bbox": [0, 210, 480, 330]},
    ]]

    boxes, page_layout = recognizer.apply_layouts([image], ocr_res, layouts, scale_factor=3)

    assert len(boxes) == 2
    assert boxes[0]["layout_type"] == "title"
    assert boxes[1]["layout_type"] == "text"
    assert len(page_layout) == 1


def test_vendored_layout_recognizer_keeps_unmatched_table_regions():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.layout_recognizer import LayoutRecognizer

    recognizer = LayoutRecognizer()
    image = np.zeros((300, 300, 3), dtype=np.uint8)
    ocr_res = [[
        {"text": "Report Title", "x0": 10, "x1": 120, "top": 10, "bottom": 25},
    ]]
    layouts = [[
        {"type": "title", "score": 0.9, "bbox": [0, 0, 360, 90]},
        {"type": "table", "score": 0.95, "bbox": [30, 120, 270, 260]},
    ]]

    boxes, _ = recognizer.apply_layouts([image], ocr_res, layouts, scale_factor=1)

    assert any(box["layout_type"] == "table" and box["text"] == "" for box in boxes)


def test_vendored_layout_recognizer_can_decode_mock_forward(monkeypatch):
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.layout_recognizer import LayoutRecognizer

    recognizer = LayoutRecognizer()
    recognizer.loaded = True
    recognizer.session = object()
    recognizer.input_name = "images"
    monkeypatch.setattr(
        recognizer,
        "_run_model_batch",
        lambda batch: [np.array([[[0.0, 0.0, 120.0, 40.0, 0.95, 2.0]]], dtype=np.float32)],
    )

    layouts = recognizer.forward([np.zeros((80, 160, 3), dtype=np.uint8)], thr=0.2, batch_size=1)

    assert len(layouts) == 1
    assert layouts[0][0]["type"] == "Title"
    assert layouts[0][0]["score"] == pytest.approx(0.95, rel=1e-6)
    assert layouts[0][0]["bbox"] == [0.0, 0.0, 30.0, 5.0]


def test_vendored_table_structure_recognizer_can_normalize_predictions():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.table_structure_recognizer import TableStructureRecognizer

    recognizer = TableStructureRecognizer()
    predictions = [[
        {"type": "table row", "score": 0.9, "bbox": [10, 10, 110, 30]},
        {"type": "table row", "score": 0.88, "bbox": [14, 30, 104, 50]},
        {"type": "table column", "score": 0.93, "bbox": [10, 10, 60, 50]},
        {"type": "table column", "score": 0.91, "bbox": [60, 12, 110, 48]},
    ]]

    normalized = recognizer(images=[np.zeros((60, 120, 3), dtype=np.uint8)], predictions=predictions)

    assert len(normalized) == 1
    assert normalized[0][0]["label"] == "table row"
    assert normalized[0][2]["label"] == "table column"
    assert normalized[0][0]["x0"] <= 14


def test_vendored_table_structure_recognizer_can_decode_mock_forward(monkeypatch):
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.table_structure_recognizer import TableStructureRecognizer

    recognizer = TableStructureRecognizer()
    recognizer.loaded = True
    recognizer.session = object()
    recognizer.input_name = "images"
    monkeypatch.setattr(
        recognizer,
        "_run_model_batch",
        lambda batch: [np.array([[[10.0, 10.0, 110.0, 30.0, 0.9, 2.0]]], dtype=np.float32)],
    )

    predictions = recognizer.forward([np.zeros((80, 160, 3), dtype=np.uint8)], thr=0.2, batch_size=1)

    assert len(predictions) == 1
    assert predictions[0][0]["type"] == "table row"


def test_vendored_table_structure_recognizer_can_construct_html_table():
    _skip_if_vision_runtime_unavailable()
    from src.shared.utils.deepdoc.vision.table_structure_recognizer import TableStructureRecognizer

    boxes = [
        {"text": "Metric", "x0": 10, "x1": 50, "top": 10, "bottom": 20, "page_number": 0, "R": "0", "C": "0", "H": True},
        {"text": "Value", "x0": 60, "x1": 100, "top": 10, "bottom": 20, "page_number": 0, "R": "0", "C": "1", "H": True},
        {"text": "Recall", "x0": 10, "x1": 50, "top": 30, "bottom": 40, "page_number": 0, "R": "1", "C": "0"},
        {"text": "0.88", "x0": 60, "x1": 100, "top": 30, "bottom": 40, "page_number": 0, "R": "1", "C": "1"},
    ]

    html = TableStructureRecognizer.construct_table(boxes, html=True)

    assert html.startswith("<table>")
    assert "<th>Metric</th>" in html
    assert "<td>Recall</td>" in html
    assert "0.88" in html


def test_vision_runtime_guard_raises_clear_error_if_parser_not_ready():
    status = get_vision_runtime_status()
    if status["parser_available"]:
        assert ensure_vision_parser_available()["parser_available"] is True
    elif status["available"]:
        with pytest.raises(DeepDocVisionParserUnavailable):
            ensure_vision_parser_available()
    else:
        with pytest.raises(DeepDocVisionRuntimeUnavailable):
            ensure_vision_parser_available()


def test_vision_runtime_guard_can_raise_runtime_error_if_dependencies_missing(monkeypatch):
    fake_report = {
        "xgboost": {"available": False},
        "pypdf": {"available": True},
        "onnxruntime": {"available": True},
        "cv2": {"available": True},
        "pillow": {"available": True},
        "huggingface_hub": {"available": True},
        "paddleocr": {"available": False},
        "fitz": {"available": True},
        "shapely": {"available": False},
        "pyclipper": {"available": False},
    }
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime.get_deepdoc_runtime_report",
        lambda: fake_report,
    )
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.vision_runtime.get_missing_runtime_dependencies",
        lambda *names: [name for name in names if not fake_report.get(name, {}).get("available")],
    )
    with pytest.raises(DeepDocVisionRuntimeUnavailable) as exc_info:
        ensure_vision_runtime_available()
    assert "xgboost" in str(exc_info.value)


def test_deepdoc_parser_supports_vision_mode():
    _skip_if_vision_runtime_unavailable()
    parser = DeepDocParser()
    result = _run(
        parser.parse_bytes(
            _build_minimal_pdf_bytes("Vision Mode"),
            file_type="pdf",
            parsing_config={"deepdoc_pdf_mode": "vision"},
            splitting_config={"chunk_size": 500},
        )
    )

    assert result.metadata["pdf_mode"] == "vision"
    assert result.metadata["vision_strategy"] == "fitz+heuristic-layout"
    assert result.metadata["layout_source"] == "heuristic"
    assert "Vision Mode" in result.full_text
    assert result.metadata["ocr_sources"] == ["fitz_text"]
    assert result.chunks


def test_ragflow_pdf_parser_vision_mode_prefers_layout_model_when_available(monkeypatch):
    _skip_if_vision_runtime_unavailable()
    parser = RAGFlowPdfParser()
    parser._layout_recognizer = SimpleNamespace(
        forward=lambda image_list, thr=0.2, batch_size=16: [[{"type": "title", "score": 0.95, "bbox": [0.0, 0.0, 1200.0, 400.0]}] for _ in image_list]
    )

    monkeypatch.setattr(
        "src.shared.utils.deepdoc.ragflow_pdf_parser.get_vision_health_status",
        lambda: {"can_run_layout_inference": True},
    )

    result = parser(_build_minimal_pdf_bytes("Model Layout"), pdf_mode="vision", chunk_size=500)

    assert result.metadata["layout_source"] == "onnx"
    assert result.metadata["layout_model_error"] is None
    assert result.metadata["vision_strategy"] == "fitz+onnx-layout"
    assert result.metadata["page_layout"][0][0]["type"] == "title"
    assert result.metadata["layout_bboxes"][0]["layout_type"] == "title"


def test_ragflow_pdf_parser_vision_mode_uses_vendored_ocr_fallback(monkeypatch):
    _skip_if_vision_runtime_unavailable()
    parser = RAGFlowPdfParser()

    monkeypatch.setattr(parser, "_extract_fitz_blocks", lambda page: [])
    monkeypatch.setattr(
        parser,
        "_extract_vendored_ocr_blocks",
        lambda image, page_index, zoom: [
            {
                "text": "OCR Block",
                "x0": 10.0,
                "x1": 80.0,
                "top": 10.0,
                "bottom": 30.0,
                "page_number": page_index,
                "font_size": 0.0,
                "ocr_source": "vendored_ocr",
            }
        ],
    )
    monkeypatch.setattr(parser, "_extract_fitz_ocr_blocks", lambda page, page_index: [])

    result = parser(_build_minimal_pdf_bytes("Invisible"), pdf_mode="vision", chunk_size=500)

    assert result.metadata["pdf_mode"] == "vision"
    assert result.metadata["ocr_sources"] == ["vendored_ocr"]
    assert "OCR Block" in result.full_text


def test_ragflow_pdf_parser_vision_mode_uses_fitz_ocr_fallback(monkeypatch):
    _skip_if_vision_runtime_unavailable()
    parser = RAGFlowPdfParser()

    monkeypatch.setattr(parser, "_extract_fitz_blocks", lambda page: [])
    monkeypatch.setattr(parser, "_extract_vendored_ocr_blocks", lambda image, page_index, zoom: [])
    monkeypatch.setattr(
        parser,
        "_extract_fitz_ocr_blocks",
        lambda page, page_index: [
            {
                "text": "Fitz OCR",
                "x0": 12.0,
                "x1": 90.0,
                "top": 15.0,
                "bottom": 35.0,
                "page_number": page_index,
                "font_size": 0.0,
                "ocr_source": "fitz_ocr",
            }
        ],
    )

    result = parser(_build_minimal_pdf_bytes("Invisible"), pdf_mode="vision", chunk_size=500)

    assert result.metadata["ocr_sources"] == ["fitz_ocr"]
    assert "Fitz OCR" in result.full_text


def test_ragflow_pdf_parser_vision_mode_preserves_table_regions_in_artifacts(monkeypatch):
    _skip_if_vision_runtime_unavailable()
    parser = RAGFlowPdfParser()
    parser._layout_recognizer = SimpleNamespace(
        forward=lambda image_list, thr=0.2, batch_size=16: [[
            {"type": "title", "score": 0.95, "bbox": [0.0, 0.0, 1200.0, 400.0]},
            {"type": "table", "score": 0.95, "bbox": [100.0, 500.0, 900.0, 1200.0]},
        ] for _ in image_list]
    )

    monkeypatch.setattr(
        "src.shared.utils.deepdoc.ragflow_pdf_parser.get_vision_health_status",
        lambda: {"can_run_layout_inference": True},
    )

    result = parser(_build_minimal_pdf_bytes("Region Table"), pdf_mode="vision", chunk_size=500)

    assert any(box["layout_type"] == "table" for box in result.metadata["layout_bboxes"])
    assert len(result.metadata["artifacts"]["tables"]) == 1
    assert result.metadata["artifacts"]["tables"][0]["pages"] == [1]
    assert len(result.metadata["table_regions"]) == 1
    assert result.metadata["table_regions"][0]["pages"] == [1]
    assert result.metadata["table_regions"][0]["page_start"] == 1
    assert result.metadata["table_regions"][0]["region_index_on_page"] == 0
    assert result.metadata["table_regions"][0]["structured_box_count"] == 0
    assert result.metadata["figure_regions"] == []
    assert any(item["kind"] == "table" for item in result.metadata["reading_order"])
    assert result.metadata["chunk_structure"]


def test_table_regions_include_page_order_and_structure_summary():
    parser = RAGFlowPdfParser()
    pdf_bytes = _build_positioned_pdf_bytes(
        [
            (72, 700, "Metric | Value"),
            (72, 684, "Recall | 0.88"),
        ]
    )

    result = parser(pdf_bytes, pdf_mode="layout", chunk_size=500)

    assert "table_regions" in result.metadata
    assert result.metadata["table_regions"] == []
    assert result.metadata["figure_regions"] == []
    assert "reading_order" in result.metadata
    assert result.metadata["reading_order"]
    assert all(item["kind"] == "text" for item in result.metadata["reading_order"])
    assert result.metadata["reading_order"][0]["global_order"] == 0
    assert result.metadata["chunk_structure"]


def test_reading_order_can_include_figure_regions():
    artifacts = {
        "tables": [],
        "figures": [
            {
                "artifact_id": "1:fig",
                "caption": "Figure 1 Trend",
                "text": "Chart region",
                "pages": [1],
                "bbox": {"x0": 20.0, "x1": 120.0, "top": 40.0, "bottom": 140.0},
                "has_image": True,
                "members": [
                    {"text": "Figure 1 Trend"},
                    {"text": "Chart region"},
                ],
            }
        ],
    }
    figure_regions = RAGFlowPdfParser._build_figure_regions_metadata(artifacts)
    reading_order = RAGFlowPdfParser._build_reading_order_metadata([], [], figure_regions)

    assert len(figure_regions) == 1
    assert figure_regions[0]["caption"] == "Figure 1 Trend"
    assert len(reading_order) == 1
    assert reading_order[0]["kind"] == "figure"
    assert reading_order[0]["global_order"] == 0


def test_parser_structured_chunks_can_include_region_blocks():
    reading_order = [
        {"kind": "text", "page": 1, "bbox": {"top": 10.0, "x0": 10.0}, "text": "Intro text", "source_id": "t1"},
        {"kind": "table", "page": 1, "bbox": {"top": 20.0, "x0": 10.0}, "caption": "Table 1", "text": "A | B", "source_id": "tb1"},
    ]

    chunks, chunk_structure = RAGFlowPdfParser._build_structured_chunks(reading_order, chunk_size=100)

    assert len(chunks) == 1
    assert "Intro text" in chunks[0]
    assert "[TABLE]" in chunks[0]
    assert chunk_structure[0]["entry_kinds"] == ["text", "table"]


def test_prepare_es_chunks_static_includes_deepdoc_chunk_metadata():
    document = SimpleNamespace(
        id=42,
        kb_id=7,
        space_id=3,
        filename="demo.pdf",
        file_type="pdf",
        file_hash="abc123",
        storage={"minio_object_name": "kb/demo.pdf"},
    )

    es_chunks = _prepare_es_chunks_static(
        document,
        ["chunk text"],
        parse_metadata={
            "parser": "deepdoc",
            "parser_class": "RAGFlowPdfParser",
            "pdf_mode": "layout",
            "layout_source": "heuristic",
            "vision_strategy": "",
            "table_regions": [{"artifact_id": "t1"}],
            "figure_regions": [{"artifact_id": "f1"}],
            "reading_order": [{"kind": "text"}, {"kind": "table"}],
            "chunk_structure": [
                {
                    "entry_kinds": ["text", "table"],
                    "entry_source_ids": ["src1", "src2"],
                    "pages": [1],
                    "entry_count": 2,
                }
            ],
        },
    )

    assert len(es_chunks) == 1
    assert es_chunks[0]["metadata"]["parser"] == "deepdoc"
    assert es_chunks[0]["metadata"]["parser_class"] == "RAGFlowPdfParser"
    assert es_chunks[0]["metadata"]["pdf_mode"] == "layout"
    assert es_chunks[0]["metadata"]["layout_source"] == "heuristic"
    assert es_chunks[0]["metadata"]["table_region_count"] == 1
    assert es_chunks[0]["metadata"]["figure_region_count"] == 1
    assert es_chunks[0]["metadata"]["reading_order_count"] == 2
    assert es_chunks[0]["metadata"]["chunk_entry_kinds"] == ["text", "table"]
    assert es_chunks[0]["metadata"]["chunk_entry_source_ids"] == ["src1", "src2"]
    assert es_chunks[0]["metadata"]["chunk_pages"] == [1]
    assert es_chunks[0]["metadata"]["chunk_entry_count"] == 2


def test_extract_parse_metadata_summary_counts_regions():
    summary = _extract_parse_metadata_summary(
        {
            "parser_class": "RAGFlowPdfParser",
            "pdf_mode": "vision",
            "layout_source": "onnx",
            "vision_strategy": "fitz+onnx-layout",
            "table_regions": [{"artifact_id": "t1"}, {"artifact_id": "t2"}],
            "figure_regions": [{"artifact_id": "f1"}],
            "reading_order": [{"kind": "text"}, {"kind": "table"}, {"kind": "figure"}],
        }
    )

    assert summary["parser_class"] == "RAGFlowPdfParser"
    assert summary["pdf_mode"] == "vision"
    assert summary["layout_source"] == "onnx"
    assert summary["vision_strategy"] == "fitz+onnx-layout"
    assert summary["table_region_count"] == 2
    assert summary["figure_region_count"] == 1
    assert summary["reading_order_count"] == 3


def test_deepdoc_engine_parse_with_parser_id_plain_pdf():
    engine = DeepDocEngine()
    result = engine.parse_with_parser_id(
        file_type="pdf",
        parser_id="pdf_plain",
        file_bytes=_build_minimal_pdf_bytes("Parser ID Plain"),
        splitting_config={"chunk_size": 500},
    )

    assert result.metadata["file_type"] == "pdf"
    assert result.metadata["pdf_mode"] == "plain"
    assert "Parser ID Plain" in result.full_text


def test_deepdoc_engine_parse_with_parser_id_excel():
    Workbook = pytest.importorskip("openpyxl").Workbook
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "KPI"
    sheet.append(["Metric", "Value"])
    sheet.append(["Recall", "0.88"])
    buffer = BytesIO()
    workbook.save(buffer)

    engine = DeepDocEngine()
    result = engine.parse_with_parser_id(
        file_type="xlsx",
        parser_id="excel",
        file_bytes=buffer.getvalue(),
        splitting_config={"chunk_size": 500},
    )

    assert result.metadata["file_type"] == "xlsx"
    assert result.metadata["parser_class"] == "RAGFlowExcelParser"
    assert "Recall" in result.full_text


def test_deepdoc_engine_exposes_runtime_dependencies():
    engine = DeepDocEngine()
    report = engine.runtime_dependencies()
    assert report["pdfplumber"]["available"] is True
    assert "xgboost" in report


def test_deepdoc_engine_exposes_vision_model_status():
    engine = DeepDocEngine()
    status = engine.vision_model_status()
    assert "groups" in status
    assert "ocr" in status["groups"]


def test_deepdoc_engine_exposes_text_concat_model_status():
    engine = DeepDocEngine()
    status = engine.text_concat_model_status()
    assert status["filename"] == "updown_concat_xgb.model"


def test_deepdoc_engine_exposes_vision_health_and_smoke_check():
    engine = DeepDocEngine()
    health = engine.vision_health_status()
    smoke = engine.vision_smoke_check()
    assert "model_groups" in health
    assert "checks" in smoke


def test_upstream_parser_package_exports_aliases():
    assert UpstreamPdfParserAlias is RAGFlowPdfParser
    assert UpstreamDocxParserAlias.__name__ == "RAGFlowDocxParser"


def test_upstream_mineru_parser_class_is_vendored():
    from src.shared.utils.deepdoc.parser.mineru_parser import (
        MinerUBackend,
        MinerUContentType,
        MinerULanguage,
        MinerUParseMethod,
        MinerUParseOptions,
        MinerUParser,
    )

    parser = MinerUParser()
    assert parser.__class__.__name__ == "MinerUParser"
    assert hasattr(parser, "parse_pdf")
    assert hasattr(parser, "check_installation")
    assert MinerUBackend.__name__ == "MinerUBackend"
    assert MinerUContentType.__name__ == "MinerUContentType"
    assert MinerULanguage.__name__ == "MinerULanguage"
    assert MinerUParseMethod.__name__ == "MinerUParseMethod"
    assert MinerUParseOptions.__name__ == "MinerUParseOptions"


def test_upstream_somark_parser_class_is_vendored():
    from src.shared.utils.deepdoc.parser.somark_parser import SoMarkAPIError, SoMarkBlockType, SoMarkParser

    parser = SoMarkParser()
    assert parser.__class__.__name__ == "SoMarkParser"
    assert hasattr(parser, "parse_pdf")
    assert hasattr(parser, "check_installation")
    assert SoMarkAPIError.__name__ == "SoMarkAPIError"
    assert SoMarkBlockType.__name__ == "SoMarkBlockType"


def test_upstream_tcadp_parser_class_is_vendored():
    from src.shared.utils.deepdoc.parser.tcadp_parser import TCADPParser, TencentCloudAPIClient

    parser = TCADPParser()
    assert parser.__class__.__name__ == "TCADPParser"
    assert hasattr(parser, "parse_pdf")
    assert hasattr(parser, "check_installation")
    assert TencentCloudAPIClient.__name__ == "TencentCloudAPIClient"


def test_upstream_resume_package_is_vendored():
    from src.shared.utils.deepdoc.parser import refactor_resume
    from src.shared.utils.deepdoc.parser.resume.step_one import FIELDS
    from src.shared.utils.deepdoc.parser.resume.step_two import highest_degree

    result = refactor_resume(
        {
            "basic": {"name": "Alice"},
            "contact": {},
            "work": [],
            "education": [],
        }
    )

    assert result["contact"]["name"] == "Alice"
    assert result["is_deleted"] == 0
    assert len(FIELDS) == 51
    assert highest_degree(["MBA"]) == "MBA"


def test_deepdoc_capabilities_expose_resume_module():
    capabilities = get_deepdoc_capabilities()

    assert capabilities["specialized_modules"]["resume"]["available"] is True
    assert capabilities["specialized_modules"]["resume"]["entrypoint"].endswith("resume.refactor")


def test_upstream_snapshot_is_exposed():
    snapshot = get_upstream_deepdoc_snapshot()
    engine_snapshot = DeepDocEngine.upstream_snapshot()

    assert snapshot["repository"] == "https://github.com/infiniflow/ragflow"
    assert snapshot["commit"]
    assert "parser" in snapshot["mirrored_packages"]
    assert snapshot == engine_snapshot


def test_deepdoc_capabilities_include_upstream_snapshot():
    capabilities = get_deepdoc_capabilities()

    assert "upstream_snapshot" in capabilities
    assert capabilities["upstream_snapshot"]["commit"]
    assert "pdf_docling" in capabilities["parser_ids"]
    assert "pdf_mineru" in capabilities["parser_ids"]
    assert "pdf_opendataloader" in capabilities["parser_ids"]
    assert "pdf_paddleocr" in capabilities["parser_ids"]
    assert "pdf_somark" in capabilities["parser_ids"]
    assert "pdf_tcadp" in capabilities["parser_ids"]


def test_deepdoc_capabilities_expose_tcadp_even_without_sdk(monkeypatch):
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.ragflow_tcadp_parser.TENCENTCLOUD_SDK_AVAILABLE",
        False,
    )
    monkeypatch.setenv("TCADP_SECRET_ID", "secret-id")
    monkeypatch.setenv("TCADP_SECRET_KEY", "secret-key")

    capabilities = get_deepdoc_capabilities()

    assert "pdf_tcadp" in capabilities["parser_ids"]
    assert capabilities["pdf_modes"]["tcadp"]["available"] is False
    assert capabilities["pdf_modes"]["tcadp"]["missing"] == ["Tencent Cloud SDK is not installed"]


def test_create_deepdoc_app_exposes_routes():
    app = create_deepdoc_app()
    routes = {route.path for route in app.routes}

    assert "/health" in routes
    assert "/capabilities" in routes
    assert "/parse-file" in routes
    assert "/parse-bytes" in routes

    health_client = TestClient(app)
    health_payload = health_client.get("/health").json()

    for route in ("/predict/dla", "/predict/ocr", "/predict/tsr"):
        if route in routes:
            assert route not in health_payload["vision_router_errors"]
        else:
            assert route in health_payload["vision_router_errors"]


def test_deepdoc_server_parse_bytes_endpoint_works():
    class _FakeEngine:
        def describe_capabilities(self):
            return {"ok": True}

        async def aparse_with_parser_id(self, **kwargs):
            return DeepDocParseResult(
                full_text="hello",
                chunks=["hello"],
                metadata={"parser": "deepdoc", "file_type": kwargs["file_type"]},
            )

    app = create_deepdoc_app(engine=_FakeEngine())
    client = TestClient(app)
    response = client.post(
        "/parse-bytes",
        json={
            "content_base64": "aGVsbG8=",
            "file_type": "txt",
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["full_text"] == "hello"
    assert payload["metadata"]["file_type"] == "txt"


def test_deepdoc_server_ocr_endpoint_rejects_invalid_operator():
    app = create_deepdoc_app()
    client = TestClient(app)
    routes = {route.path for route in app.routes}
    if "/predict/ocr" not in routes:
        pytest.skip("OCR endpoint is not mounted in the current runtime")
    response = client.post(
        "/predict/ocr",
        files={"request": ("sample.png", b"123", "image/png")},
        data={"operator": "bad"},
    )

    assert response.status_code == 400
    assert "must be 'det' or 'rec'" in response.json()["detail"]


def test_deepdoc_engine_wraps_vision_model_download(monkeypatch, tmp_path):
    expected_path = tmp_path / "deepdoc-models"
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.engine.download_model_group",
        lambda group=None: expected_path,
    )
    engine = DeepDocEngine()
    assert engine.download_vision_models("ocr") == expected_path


def test_deepdoc_engine_wraps_vision_model_group_check(monkeypatch, tmp_path):
    monkeypatch.setattr(
        "src.shared.utils.deepdoc.engine.ensure_model_group_available",
        lambda group: tmp_path,
    )
    engine = DeepDocEngine()
    assert engine.ensure_vision_model_group("layout") == tmp_path


def test_deepdoc_result_to_documents():
    result = DeepDocParseResult(
        full_text="alpha\nbeta",
        chunks=["alpha", "beta"],
        metadata={"parser": "deepdoc"},
    )

    docs = result.to_documents(source="demo")
    assert len(docs) == 2
    assert docs[0]["chunk_index"] == 0
    assert docs[0]["source"] == "demo"
    assert docs[1]["content"] == "beta"
