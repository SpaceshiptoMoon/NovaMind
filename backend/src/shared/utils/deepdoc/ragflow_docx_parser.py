from __future__ import annotations

# Adapted from RAGFlow deepdoc/parser/docx_parser.py

import logging
import re
from io import BytesIO
from typing import Any, List, Sequence

from docx import Document
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)

from src.shared.utils.deepdoc.compat import LazyImage, rag_tokenizer
from src.shared.utils.deepdoc.constants import MAXIMUM_PAGE_NUMBER


class RAGFlowDocxParser:
    def get_picture(self, document: Any, paragraph: Any) -> LazyImage | None:
        imgs = paragraph._element.xpath(".//pic:pic")
        if not imgs:
            return None

        image_blobs: List[bytes] = []
        for img in imgs:
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                continue
            rel_id = embed[0]
            image_blob = None
            try:
                related_part = document.part.related_parts[rel_id]
            except Exception as exc:
                logging.warning("Skipping image due to unexpected related part error: %s", exc)
                continue

            try:
                image = related_part.image
                if image is not None:
                    image_blob = image.blob
            except (
                UnrecognizedImageError,
                UnexpectedEndOfFileError,
                InvalidImageStreamError,
                UnicodeDecodeError,
            ) as exc:
                logging.info("Damaged image encountered, attempting blob fallback: %s", exc)
            except Exception as exc:
                logging.warning("Unexpected error getting image, attempting blob fallback: %s", exc)

            if image_blob is None:
                image_blob = getattr(related_part, "blob", None)
            if image_blob:
                image_blobs.append(image_blob)

        if not image_blobs:
            return None
        return LazyImage(image_blobs)

    def _extract_table_content(self, table: Any) -> List[str]:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        return self._compose_table_content(rows)

    def _compose_table_content(self, rows: Sequence[Sequence[str]]) -> List[str]:
        def block_type(block: str) -> str:
            patterns = [
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日?$", "Dt"),
                (r"^(20|19)[0-9]{2}年?$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月?$", "Dt"),
                (r"^[0-9]{1,2}[月/-][0-9]{1,2}日?$", "Dt"),
                (r"^第?[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年?[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                (r"^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg"),
            ]
            for pattern, name in patterns:
                if re.search(pattern, block):
                    return name

            tokens = [t for t in rag_tokenizer.tokenize(block).split() if len(t) > 1]
            if len(tokens) > 3:
                return "Tx" if len(tokens) < 12 else "Lx"
            if len(tokens) == 1 and rag_tokenizer.tag(tokens[0]) == "nr":
                return "Nr"
            return "Ot"

        if len(rows) < 2:
            return []

        column_count = max(len(row) for row in rows)
        normalized_rows = [list(row) + [""] * (column_count - len(row)) for row in rows]
        data_rows = normalized_rows[1:]

        dominant_counts = {}
        for row in data_rows:
            for cell in row:
                name = block_type(str(cell))
                dominant_counts[name] = dominant_counts.get(name, 0) + 1
        dominant_type = max(dominant_counts.items(), key=lambda item: item[1])[0]

        header_rows = [0]
        if dominant_type == "Nu":
            for row_index in range(1, len(normalized_rows)):
                row_types = {}
                for cell in normalized_rows[row_index]:
                    name = block_type(str(cell))
                    row_types[name] = row_types.get(name, 0) + 1
                row_type = max(row_types.items(), key=lambda item: item[1])[0]
                if row_type != dominant_type:
                    header_rows.append(row_index)

        lines: List[str] = []
        for row_index in range(1, len(normalized_rows)):
            if row_index in header_rows:
                continue

            related_headers = [idx for idx in header_rows if idx < row_index]
            if len(related_headers) > 1:
                trimmed = [related_headers[-1]]
                for i in range(len(related_headers) - 2, -1, -1):
                    if trimmed[0] - related_headers[i] <= 1:
                        trimmed.insert(0, related_headers[i])
                    else:
                        break
                related_headers = trimmed

            headers: List[str] = []
            for column_index in range(column_count):
                pieces: List[str] = []
                for header_index in related_headers:
                    value = str(normalized_rows[header_index][column_index]).strip()
                    if value and value not in pieces:
                        pieces.append(value)
                header = ",".join(pieces)
                headers.append(f"{header}: " if header else "")

            cells: List[str] = []
            for column_index in range(column_count):
                value = str(normalized_rows[row_index][column_index]).strip()
                if value:
                    cells.append(headers[column_index] + value)
            lines.append(";".join(cells))

        if column_count > 3:
            return lines
        return ["\n".join(lines)] if lines else []

    def __call__(self, file_name_or_binary: str | bytes, from_page: int = 0, to_page: int = MAXIMUM_PAGE_NUMBER):
        self.doc = (
            Document(file_name_or_binary)
            if isinstance(file_name_or_binary, str)
            else Document(BytesIO(file_name_or_binary))
        )
        parsed_page = 0
        sections = []

        for paragraph in self.doc.paragraphs:
            if parsed_page > to_page:
                break

            runs_within_single_paragraph: List[str] = []
            for run in paragraph.runs:
                if parsed_page > to_page:
                    break
                if from_page <= parsed_page < to_page and paragraph.text.strip():
                    runs_within_single_paragraph.append(run.text)

                if "lastRenderedPageBreak" in run._element.xml:
                    parsed_page += 1

            style_name = paragraph.style.name if hasattr(paragraph.style, "name") else ""
            image = self.get_picture(self.doc, paragraph)
            sections.append(
                {
                    "text": "".join(runs_within_single_paragraph),
                    "style": style_name,
                    "image": image,
                }
            )

        tables = [self._extract_table_content(table) for table in self.doc.tables]
        return sections, tables
