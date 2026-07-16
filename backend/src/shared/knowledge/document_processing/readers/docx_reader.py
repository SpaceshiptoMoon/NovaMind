import os
from typing import List, Dict
from novamind.shared.knowledge.document_processing.readers.base_reader import BaseReader
from novamind.shared.knowledge.document_processing.readers.executor import run_in_executor
from novamind.core.middleware.structured_logging import get_logger
from docx import Document

logger = get_logger(__name__)


class DocxReader(BaseReader):
    """DOCX文档读取器"""

    def __init__(self):
        super().__init__()

    def _load_data_sync(self, file_path: str) -> List[Dict[str, str]]:
        """
        同步读取 DOCX 文件（在线程池中执行）
        :param file_path: DOCX文件路径
        :return: 文档块列表
        """
        documents = []
        try:
            # 使用python-docx读取DOCX
            doc = Document(file_path)

            # 提取所有段落的文本
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            # 获取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            text = '\n'.join(full_text)

            if text.strip():
                documents.append({
                    'text': text,
                    'source': os.path.basename(file_path),
                    'page': 1,
                    'doc_id': f"docx_{hash(file_path)}",
                    'type': 'docx'
                })
        except Exception as e:
            logger.error("读取DOCX文件失败", file_path=str(file_path), error=str(e))

        return documents

    async def load_data(self, file_path: str) -> List[Dict[str, str]]:
        """
        从DOCX文件加载数据（异步，在共享线程池中执行）
        :param file_path: DOCX文件路径
        :return: 文档块列表，每个文档块包含文本和其他元数据
        """
        return await run_in_executor(self._load_data_sync, file_path)